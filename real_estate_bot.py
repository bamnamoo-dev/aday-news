import os
import requests
from bs4 import BeautifulSoup
from datetime import datetime, timedelta
import json
import time

import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.header import Header
from dotenv import load_dotenv

# 로컬 .env 파일 로드 (환경 변수 설정)
load_dotenv()

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    pass

# ==========================================
# 1. 설정 (Settings)
# ==========================================
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "YOUR_DEFAULT_OR_TEST_KEY")
SAVE_DIR = os.getenv("SAVE_DIR", "reports")
MAX_NEWS_COUNT = 10

# 이메일 설정
EMAIL_SENDER = os.getenv("EMAIL_USER")       # 보내는 사람 (본인 지메일)
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD") # 구글 앱 비밀번호
EMAIL_RECEIVER = "bamnamoo@gmail.com"        # 받는 사람

# 키워드 설정 (부동산 핵심 키워드)
INCLUDE_KEYWORDS = [
    '부동산', '아파트', '집값', '분양', '재건축', '재개발', '전세', '월세', 
    '매매', '주택', 'LH', 'SH', '청약', '임대', '건설', '공급', '규제', 
    '대출', '금리', '국토부', '빌라', '오피스텔', '단지', 'GTX', '공공주택',
    '주거', '경매', '공시지가', '매수', '매도', '수도권', '시장동향'
]
# 제외 키워드 (부동산과 무관한 노이즈 제거)
EXCLUDE_KEYWORDS = [
    '사고', '화재', '경찰', '살인', '폭행', '충돌', '추락', '폭발', '사망', 
    '부상', '피해', '혐의', '구속', '검거', '사기', '횡령', '음주', '마약',
    '정전', '구조', '급락', '증시', '삼성전자', '코스피', '코스닥', '공모주',
    '나스닥', '뉴욕증시', '환율', '반도체', '배터리', '자동차', '전기차',
    '아이폰', '갤럭시', '물가', '금통위', '한은', '가계부채', '자영업', 
    '소상공인', '편의점', '식당', '카페', '프랜차이즈', '최저임금', '노동계', 
    '알바', '영업이익', '매출', '소비자', '물가상승', '인건비', '삼전', '주식', '주가'
]

# ==========================================
# 2. 뉴스 수집 로봇 (Crawler)
# ==========================================
def get_ranking_news(date_str):
    # 변경: 경제 내 '부동산' 섹션 랭킹 (이게 가장 확실합니다)
url = f"https://news.naver.com/main/ranking/popularDay.naver?sectionId=101&subSectionId=261&date={date_str}"
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"}
    
    try:
        response = requests.get(url, headers=headers)
        soup = BeautifulSoup(response.text, 'html.parser')
        news_list = []
        
        for office_list in soup.select('.rankingnews_box'):
            for item in office_list.select('.rankingnews_list > li'):
                title_tag = item.select_one('a')
                if not title_tag: continue
                
                title = title_tag.get_text(strip=True)
                link = title_tag['href']
                
                if link.startswith('/'):
                    link = "https://news.naver.com" + link
                
                is_real_estate = any(k in title for k in INCLUDE_KEYWORDS)
                is_excluded = any(k in title for k in EXCLUDE_KEYWORDS)
                
                if is_real_estate and not is_excluded:
                    news_list.append({"title": title, "link": link})
        
        return news_list
    except Exception as e:
        print(f"[{date_str}] 뉴스 목록 수집 중 오류: {e}")
        return []

def get_article_content(url):
    """뉴스 원문 내용을 추출"""
    headers = {"User-Agent": "Mozilla/5.0"}
    try:
        response = requests.get(url, headers=headers)
        soup = BeautifulSoup(response.text, 'html.parser')
        
        content_tag = soup.select_one('#dic_area') or soup.select_one('#articleBodyContents')
        if content_tag:
            for s in content_tag.select('script, style, .article_footer, .img_desc'):
                s.decompose()
            return content_tag.get_text(separator="\n", strip=True)
        return "본문을 불러올 수 없습니다."
    except Exception as e:
        return f"원문 추출 오류: {e}"

# ==========================================
# 3. AI 분석 로봇 (Gemini 3)
# ==========================================
def analyze_with_gemini(news_data):
    """Gemini API를 사용하여 블로그 원고와 요약 리스트를 생성"""
    print("AI가 블로그 포스팅 원고와 요약을 작성 중입니다...")
    
    context_text = ""
    for i, item in enumerate(news_data):
        content_snippet = item['content'][:2000] 
        context_text += f"\n[뉴스 {i+1}]\n제목: {item['title']}\n링크: {item['link']}\n내용: {content_snippet}\n"
    
    # 프롬프트 부분 수정 예시
prompt = f"""
당신은 부동산 전문 분석가입니다. 아래 제공된 뉴스 목록 중 **부동산 시장 흐름, 정책, 투자, 분양**과 직접적인 관련이 있는 뉴스만 선별하여 분석하세요.

[필터링 규칙]
- 단순 화재, 정전, 고독사, 살인사건 등 '사회적 사건/사고'는 부동산 가치 분석에 불필요하므로 절대 포함하지 마세요.
- 연예인 집 공개, 단순 가십성 뉴스도 제외하세요.
- 오직 시장 전망, 정책 변화, 신도시 소식, 재개발 현황 등 '경제적 가치'가 있는 뉴스만 최대 10개 선택하세요.

    [작성 항목]
    1. 오늘의 주요 뉴스 및 요약 링크: 
       각 뉴스(1~10)에 대해 아래 형식을 엄격히 지켜 작성하세요.
       형식:
       [번호]. [뉴스 제목]
       🔗 [뉴스 링크]
       📝 요약: [해당 뉴스의 핵심 내용을 정확히 2줄로 요약]

    2. 블로그 포스팅 원고: 독자들이 흥미를 가질만한 제목, 서론, 본문(흐름 분석 및 전략), 결론, 해시태그를 포함한 완성된 블로그 원고.

    구분선인 '---BLOG_START---' 를 사용하여 1번 항목과 2번 항목을 구분해 주세요.

    [입력 데이터]
    {context_text}
    """

    payload = {
        "contents": [{"parts": [{"text": prompt}]}]
    }
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3-flash-preview:generateContent?key={GEMINI_API_KEY}"
    
    try:
        response = requests.post(url, json=payload, headers={'Content-Type': 'application/json'})
        if response.status_code == 200:
            return response.json()['candidates'][0]['content']['parts'][0]['text']
        else:
            return f"AI 분석 오류: {response.text}"
    except Exception as e:
        return f"AI 통신 오류: {e}"

def save_as_docx(full_path, report_content):
    """보고서 내용을 .docx 파일로 저장"""
    try:
        doc = Document()
        title_line = report_content.split('\n')[0]
        doc.add_heading(title_line, 0)
        for line in report_content.split('\n')[1:]:
            if line.startswith('='):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            elif line.startswith('###'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.startswith('####'):
                doc.add_heading(line.replace('#', '').strip(), level=2)
            else:
                doc.add_paragraph(line)
        doc.save(full_path)
        return True
    except Exception as e:
        print(f"DOCX 저장 중 오류: {e}")
        return False

def send_email(subject, body):
    """분석 리포트를 이메일로 발송"""
    if not EMAIL_SENDER or not EMAIL_PASSWORD:
        print("공지: 이메일 설정(EMAIL_USER, EMAIL_PASSWORD)이 되어있지 않아 메일 발송을 건너뜁니다.")
        return False

    try:
        msg = MIMEMultipart()
        msg['From'] = EMAIL_SENDER
        msg['To'] = EMAIL_RECEIVER
        msg['Subject'] = Header(subject, 'utf-8')
        
        msg.attach(MIMEText(body, 'plain', 'utf-8'))
        
        # SMTP 서버 연결 (Gmail 기준)
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(EMAIL_SENDER, EMAIL_PASSWORD)
            server.sendmail(EMAIL_SENDER, EMAIL_RECEIVER, msg.as_string())
        
        print(f"성공! 이메일이 발송되었습니다: {EMAIL_RECEIVER}")
        return True
    except Exception as e:
        print(f"이메일 발송 중 오류: {e}")
        return False

# ==========================================
# 4. 메인 실행 (Main)
# ==========================================
def main():
    print("부동산 뉴스 자동 분석 프로그램을 시작합니다.")
    
    today = datetime.now()
    yesterday = today - timedelta(days=1)
    dates = [today.strftime('%Y%m%d'), yesterday.strftime('%Y%m%d')]
    
    collected_news = []
    seen_links = set()
    
    for date_str in dates:
        print(f"날짜 {date_str} 경제 랭킹 뉴스를 수집 중...")
        news_items = get_ranking_news(date_str)
        
        for item in news_items:
            if item['link'] not in seen_links:
                seen_links.add(item['link'])
                collected_news.append(item)
            
            if len(collected_news) >= MAX_NEWS_COUNT:
                break
        if len(collected_news) >= MAX_NEWS_COUNT:
            break

    if not collected_news:
        print("분석할 부동산 뉴스를 찾지 못했습니다.")
        return

    print(f"총 {len(collected_news)}개의 뉴스 원문을 추출합니다...")
    for item in collected_news:
        item['content'] = get_article_content(item['link'])
        time.sleep(0.5)

    # AI 분석 (요약 + 블로그 스타일)
    ai_output = analyze_with_gemini(collected_news)
    
    # 요약과 블로그 본문 분리
    if "---BLOG_START---" in ai_output:
        summaries_part, blog_post = ai_output.split("---BLOG_START---")
    else:
        summaries_part = "요약을 생성하지 못했습니다."
        blog_post = ai_output

    # 최종 리포트 조립
    print("블로그 스타일로 최종 보고서를 생성 중입니다...")
    report = f"📅 생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    report += "🚀 본 파일은 블로그 포스팅에 바로 사용할 수 있도록 구성되었습니다.\n\n"
    
    report += "="*60 + "\n"
    report += "          [ 1단계: 오늘의 주요 뉴스 및 요약 링크 ]\n"
    report += "="*60 + "\n\n"
    
    # AI가 생성한 뉴스별 [제목/링크/요약] 리스트 배치
    report += summaries_part.strip() + "\n"
    
    report += "\n" + "="*60 + "\n"
    report += "          [ 2단계: 블로그 포스팅 원고 ]\n"
    report += "="*60 + "\n\n"
    report += blog_post.strip()
    
    report += "\n\n" + "-"*60 + "\n"
    report += "제작: 니크의 부동산 정보\n"

    # 저장
    if not os.path.exists(SAVE_DIR):
        os.makedirs(SAVE_DIR)
    
    base_name = f"오늘의부동산포스트_{today.strftime('%Y%m%d_%H%M%S')}"
    
    # 1. TXT 저장
    txt_path = os.path.join(SAVE_DIR, base_name + ".txt")
    try:
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(report)
        print(f"\n성공! TXT 보고서가 저장되었습니다: {txt_path}")
    except Exception as e:
        print(f"TXT 저장 중 오류: {e}")
        
    # 2. DOCX 저장
    docx_path = os.path.join(SAVE_DIR, base_name + ".docx")
    if 'Document' in globals():
        if save_as_docx(docx_path, report):
            print(f"성공! DOCX 보고서가 저장되었습니다: {docx_path}")
    else:
        print("공지: 'python-docx' 라이브력이 없어 .docx 저장을 건너뜁니다.")

    # 3. 이메일 발송
    email_subject = f"[오늘의 부동산 포스트] {today.strftime('%Y-%m-%d')} 분석 리포트"
    send_email(email_subject, report)

if __name__ == "__main__":
    main()
